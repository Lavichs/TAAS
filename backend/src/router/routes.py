import datetime
import io
import uuid

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.shared import Pt, Inches, RGBColor
from flask import Blueprint, jsonify, request, send_file
from flask_cors import CORS

from src.db.Models import *
from utils.dictionaries import check_required_fields
from utils.strings import randomword

main = Blueprint('main', __name__)
CORS(main)
LP_tokens = dict()
tokens_info = dict()


@main.route('/tours')
def get_tours():
    tours = db.session.query(Tour, TourOperator).join(TourOperator).filter(Tour.isDelete == False).all()
    data = []
    for tour, to in tours:
        data.append({
            'id': tour.id,
            'country': tour.country_stay,
            'city1': tour.city_departure,
            'city2': tour.city_stay,
            'date1': tour.date_departure,
            'date2': tour.date_return,
            'description': tour.description,
            'price': tour.price,
            'tourOperator': to.title
        })
    customHeaders = {
        "xTotalCount": len(tours)
    }
    response = jsonify([customHeaders, data])
    response.status_code = 200
    response.headers.add('Content-Type', 'application/json; charset=utf-8')
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('xTotalCount', str(len(data)))
    return response


@main.route('/tour/c', methods=['POST'])
def create_tour():
    data = request.json
    if not data:
        response = jsonify({'error': 'Invalid data provided'})
        response.status_code = 400
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
    new_tour = Tour(
        id=uuid.uuid4(),
        country_stay=data.get('country'),
        city_departure=data.get('city1'),
        city_stay=data.get('city2'),
        date_departure=datetime.datetime.strptime(data.get('date1'), '%Y-%m-%d'),
        date_return=datetime.datetime.strptime(data.get('date2'), '%Y-%m-%d'),
        description=data.get('description'),
        price=data.get('price'),
        idTourOperator=uuid.UUID(data.get('tourOperator')),
        isDelete=False
    )
    db.session.add(new_tour)
    db.session.commit()
    print(new_tour)
    response = jsonify({'message': 'Tour created successfully', 'id': new_tour.id})
    # response = jsonify({'message': 'Successfully'})
    response.status_code = 201
    response.headers.add('Access-Control-Allow-Origin', '*')
    return response


@main.route('/bookings')
def getBookings():
    limit = request.args.get('_limit')
    page = request.args.get('_page')
    # bookings = db.session.query(Tour, TourOperator).join(TourOperator).all()
    # bookings = db.session.query(Client, Employee).join(TourOperator).all()
    bookings = db.session.query(Booking, Client, Employee, Tour) \
        .join(Client, Booking.idClient == Client.id) \
        .join(Employee, Booking.idEmployee == Employee.id) \
        .join(Tour, Booking.idTour == Tour.id) \
        .all()
    bookings2 = db.session.query(Booking).all()
    data = []
    for booking, client, employee, tour in bookings:
        data.append({
            "id": booking.id,
            "surnameClient": client.surname,
            "nameClient": client.name,
            "patronymicClient": client.patronymic,
            "surnameEmployee": employee.surname,
            "nameEmployee": employee.name,
            "patronymicEmployee": employee.patronymic,
            "cityTo": tour.city_stay,
            "cityFrom": tour.city_departure,
            "date1": tour.date_departure,
            "date2": tour.date_return,
            "status": booking.status,
            "tourId": tour.id,
            "pay_method": booking.pay_method,
            "idClient": booking.idClient,
            "cost": booking.cost,
            "country": tour.country_stay
        })
    customHeaders = {
        "xTotalCount": len(bookings)
    }
    response = jsonify([customHeaders, data])
    response.status_code = 200
    response.headers.add('Content-Type', 'application/json; charset=utf-8')
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('xTotalCount', str(len(data)))
    return response


@main.route('/login', methods=['POST'])
def login():
    data = request.json
    missing_fields = check_required_fields(data, ['login', 'password'])
    if missing_fields:
        response = jsonify({'error': 'Invalid data provided'})
        response.status_code = 400
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
    user = db.session.query(Employee) \
        .filter_by(login=data.get('login')) \
        .filter_by(encrypted_password=data.get('password')) \
        .all()
    if not user:
        response = jsonify()
        response.status_code = 401
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
    token_key = ''.join([data.get('login'), data.get('password')])
    if LP_tokens.get(token_key) is None:
        new_token = randomword(20)
        LP_tokens[token_key] = new_token
        tokens_info[new_token] = user[0].id
    else:
        new_token = LP_tokens.get(token_key)
    response = jsonify({'token': new_token, 'role': user[0].post})
    response.status_code = 200
    response.headers.add('Access-Control-Allow-Origin', '*')
    return response


@main.route('/report')
def generate_docx():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    # Создание нового стиля
    new_style = doc.styles.add_style('MyMainHeadingStyle', 1)  # 1 - это уровень заголовка
    new_style.font.name = 'Times New Roman'
    new_style.font.size = Pt(16)
    new_style.font.color.rgb = RGBColor(0, 0, 0)
    new_style.font.bold = True
    # Получаем текущую дату
    current_date = datetime.datetime.now()
    # Получаем первый день текущего месяца
    first_day_of_month = current_date.replace(day=1)
    # Получаем последний день текущего месяца
    last_day_of_month = current_date.replace(day=1, month=current_date.month % 12 + 1) - datetime.timedelta(days=1)
    # Получаем название месяца и год
    month_name = current_date.strftime('%B')
    year = current_date.year
    pathToImg = 'resources/albatros.jpg'
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    paragraph_icon = doc.add_paragraph()
    run = paragraph_icon.add_run()
    run.add_picture(pathToImg, width=Inches(1.0))
    paragraph_icon.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_icon.paragraph_format.space_after = 0
    paragraph_about = doc.add_paragraph(
        f'г. Оренбург\n{".".join((str(datetime.datetime.now().date()).split("-")[::-1]))}')
    paragraph_about.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph_about.paragraph_format.space_after = 0
    heading = doc.add_heading('Отчёт о продажах', level=1)
    heading.style = doc.styles['MyMainHeadingStyle']
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading.paragraph_format.space_before = 0
    heading.paragraph_format.space_after = 0
    paragraph_predislovie = doc.add_paragraph(
        '\nОбщество с ограниченной ответственностью "Альбатрос"'
        # ', в дальнейшем именуемое "Предприятие", '
        f'в период с {str(first_day_of_month).split(" ")[0]} по {str(last_day_of_month).split(" ")[0]}'
        f'реализовало следующие товары*:\n')
    paragraph_predislovie.paragraph_format.space_after = 0
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    paragraph_c = hdr_cells[0].paragraphs[0]
    run = paragraph_c.add_run('№ п/п')
    run.bold = True
    paragraph_c = hdr_cells[1].paragraphs[0]
    run = paragraph_c.add_run('Id')
    run.bold = True
    paragraph_c = hdr_cells[2].paragraphs[0]
    run = paragraph_c.add_run('Количество продаж')
    run.bold = True
    paragraph_c = hdr_cells[3].paragraphs[0]
    run = paragraph_c.add_run('Цена ₽')
    run.bold = True
    total_count, total_sum = 0, 0
    p_p = 1
    # tour_count = self.api_native.getCountTourIdInBooking()
    bookings = db.session.query(Booking) \
        .join(Tour, Booking.idTour == Tour.id) \
        .all()
    bookingIds = dict()
    for booking in bookings:
        print(booking)
        key = booking.idTour
        if key in bookingIds:
            bookingIds[key][0] += 1
            bookingIds[key][2] += booking.cost
        else:
            bookingIds[key] = [1, booking.cost, booking.cost, booking]
    for key, value in bookingIds.items():
        count, price, total_price, booking = value
        row_cells = table.add_row().cells
        row_cells[0].text = str(p_p)
        row_cells[1].text = str(key)[:8]
        row_cells[2].text = str(count)
        row_cells[3].text = str(price)
        total_count += count
        total_sum += total_price
        p_p += 1
    row_cells = table.add_row().cells
    paragraph_c = row_cells[2].paragraphs[0]
    run = paragraph_c.add_run(str(total_count))
    run.bold = True
    paragraph_c = row_cells[3].paragraphs[0]
    run = paragraph_c.add_run(str(total_sum))
    run.bold = True
    row_cells[0].merge(row_cells[1])
    paragraph_c = row_cells[0].paragraphs[0]
    run = paragraph_c.add_run('Итого: ')
    run.bold = True
    paragraph_PS = doc.add_paragraph('\n* - описание товаров соответственно Id приведено в Приложении 1')
    paragraph_PS.paragraph_format.space_after = 0
    paragraph_end = doc.add_paragraph('\n\nГлавный бухгалтер:\t______________\t\t\t______________')
    paragraph_end.paragraph_format.space_after = 0
    # Добавляем разрыв страницы
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    heading_P_1 = doc.add_heading(f'Приложение 1\n', level=1)
    heading_P_1.style = doc.styles['MyMainHeadingStyle']
    heading_P_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    heading_P_1.paragraph_format.space_before = 0
    heading_P_1.paragraph_format.space_after = 0
    table_1 = doc.add_table(rows=1, cols=6)
    hdr_cells_1 = table_1.rows[0].cells
    paragraph_pc = hdr_cells_1[0].paragraphs[0]
    run = paragraph_pc.add_run('Id')
    run.bold = True
    paragraph_pc = hdr_cells_1[1].paragraphs[0]
    run = paragraph_pc.add_run('Страна')
    run.bold = True
    paragraph_pc = hdr_cells_1[2].paragraphs[0]
    run = paragraph_pc.add_run('Пункт отправления')
    run.bold = True
    paragraph_pc = hdr_cells_1[3].paragraphs[0]
    run = paragraph_pc.add_run('Пункт назначения')
    run.bold = True
    paragraph_pc = hdr_cells_1[4].paragraphs[0]
    run = paragraph_pc.add_run('Дата отправления')
    run.bold = True
    paragraph_pc = hdr_cells_1[5].paragraphs[0]
    run = paragraph_pc.add_run('Дата назначения')
    run.bold = True
    tours = dict()
    for bId, v in bookingIds.items():
        tour = Tour.query.filter_by(id=bId).first()
        tours[tour.id] = tour
    for key, value in tours.items():
        row_cells = table_1.add_row().cells
        row_cells[0].text = str(key)[:8]
        row_cells[1].text = str(value.country_stay)
        row_cells[2].text = str(value.city_departure)
        row_cells[3].text = str(value.city_stay)
        row_cells[4].text = str(value.date_departure.date())
        row_cells[5].text = str(value.date_return.date())
    temp_file = io.BytesIO()
    doc.save(temp_file)
    temp_file.seek(0)
    #
    return send_file(temp_file, as_attachment=True, download_name='generated_file.docx')


@main.route('/create/booking', methods=['POST'])
def create_booking():
    data = request.json
    missing_fields = check_required_fields(data, ['token', 'email', 'surname', 'name', 'patronymic', 'tel',
                                                  'seriesPassport', 'numberPassport', 'tourId', 'tourPrice'])
    if missing_fields:
        response = jsonify({'error': 'Invalid data provided'})
        response.status_code = 400
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
    existing_contact = Contact.query.filter_by(
        email=data.get('email'),
        tel_number=data.get('tel')) \
        .first()
    if existing_contact:
        contact = existing_contact
    else:
        contact = Contact(
            id=uuid.uuid4(),
            email=data.get('email'),
            tel_number=data.get('tel')
        )
        db.session.add(contact)
        db.session.commit()
    existing_client = Client.query.filter_by(
        name=data.get('name'),
        surname=data.get('surname'),
        patronymic=data.get('patronymic'),
        pasport_series=data.get('seriesPassport'),
        pasport_number=data.get('numberPassport'),
        contactID=contact.id) \
        .first()
    if existing_client:
        client = existing_client
    else:
        client = Client(
            id=uuid.uuid4(),
            name=data.get('name'),
            surname=data.get('surname'),
            patronymic=data.get('patronymic'),
            pasport_series=data.get('seriesPassport'),
            pasport_number=data.get('numberPassport'),
            contactID=contact.id)
        db.session.add(client)
        db.session.commit()
    booking = Booking(
        id=uuid.uuid4(),
        idClient=client.id,
        idTour=uuid.UUID(data.get('tourId')),
        idEmployee=tokens_info.get(data.get('token')),
        status='не оплачено',
        pay_method='-',
        cost=data.get('tourPrice')
    )
    db.session.add(booking)
    db.session.commit()
    response = jsonify('created')
    response.status_code = 200
    response.headers.add('Access-Control-Allow-Origin', '*')
    return response


@main.route('/update/booking', methods=['POST'])
def update_booking():
    data = request.json
    missing_fields = check_required_fields(data, ['status', 'payMethod'])
    if missing_fields:
        response = jsonify({'error': 'Invalid data provided'})
        response.status_code = 400
        response.headers.add('Access-Control-Allow-Origin', '*')
        return response
    booking = Booking.query.filter_by(id=uuid.UUID(data.get('id'))).first()
    booking.status = data.get('status')
    booking.pay_method = data.get('payMethod')
    booking.date_pay = datetime.datetime.today()
    db.session.commit()
    response = jsonify('updated')
    response.status_code = 200
    response.headers.add('Access-Control-Allow-Origin', '*')
    return response


@main.route('/delete/tour', methods=['POST'])
def remove_tour():
    data = request.json
    tour = Tour.query.filter_by(id=uuid.UUID(data.get('id'))).first()
    tour.isDelete = True
    db.session.commit()
    response = jsonify('deleted')
    response.status_code = 200
    response.headers.add('Access-Control-Allow-Origin', '*')
    return response


@main.route('/tourOperators', methods=['GET'])
def get_to():
    to = TourOperator.query.all()
    data = []
    for operator in to:
        data.append({
            'id': operator.id,
            'title': operator.title
        })
    customHeaders = {
        "xTotalCount": len(to)
    }
    response = jsonify([customHeaders, data])
    response.status_code = 200
    response.headers.add('Content-Type', 'application/json; charset=utf-8')
    response.headers.add('Access-Control-Allow-Origin', '*')
    response.headers.add('xTotalCount', str(len(data)))
    return response


@main.route('/contract', methods=['GET'])
def get_contract():
    id_booking = request.args.get('_id')
    print(id_booking)
    booking = Booking.query.filter_by(id=uuid.UUID(id_booking)).first()
    tour = Tour.query.filter_by(id=booking.idTour).first()
    client = Client.query.filter_by(id=booking.idClient).first()
    to = TourOperator.query.filter_by(id=tour.idTourOperator).first()
    cotactTO = Contact.query.filter_by(id=to.contactID).first()
    replacements = {
        '[datePrint]': f'\n{booking.date_pay.date()}\n',
        '[clientFIO]': f'{client.surname} {client.name} {client.patronymic}',
        '[bookingId]': f'{booking.id}',
        '[country]': f'{tour.country_stay}',
        '[city1]': f'{tour.city_departure}',
        '[city2]': f'{tour.city_stay}',
        '[date1]': f'{tour.date_departure.date()}',
        '[date2]': f'{tour.date_return.date()}',
        '[price]': f'{tour.price}₽',
        '[titleTO]': f'{to.title}',
        '[emailTO]': f'{cotactTO.email}',
        '[telTO]': f'{cotactTO.tel_number}',
    }
    # Открываем исходный документ
    doc = Document('resources/pattern_1.docx')
    for p in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)
                
    temp_file = io.BytesIO()
    doc.save(temp_file)
    temp_file.seek(0)
    return send_file(temp_file, as_attachment=True, download_name='generated_file.docx')

# @main.route('/contract_old', methods=['GET'])
# def get_contract_old():
#     id_booking = request.args.get('_id')
#     print(id_booking)
# 
#     booking = Booking.query.filter_by(id=uuid.UUID(id_booking)).first()
#     tour = Tour.query.filter_by(id=booking.idTour).first()
#     client = Client.query.filter_by(id=booking.idClient).first()
#     to = TourOperator.query.filter_by(id=tour.idTourOperator).first()
#     cotactTO = Contact.query.filter_by(id=to.contactID).first()
# 
#     print(tour.id)
#     print(tour.country_stay)
# 
#     doc = Document()
#     style = doc.styles['Normal']
#     font = style.font
#     font.name = 'Times New Roman'
#     font.size = Pt(14)
# 
#     # Создание нового стиля
#     new_style = doc.styles.add_style('MyMainHeadingStyle', 1)  # 1 - это уровень заголовка
#     new_style.font.name = 'Times New Roman'
#     new_style.font.size = Pt(16)
#     new_style.font.color.rgb = RGBColor(0, 0, 0)
#     new_style.font.bold = True
#     # Создание нового стиля
#     new_style = doc.styles.add_style('MyHeadingStyle', 1)  # 1 - это уровень заголовка
#     new_style.font.name = 'Times New Roman'
#     new_style.font.size = Pt(14)
#     new_style.font.color.rgb = RGBColor(0, 0, 0)
#     new_style.font.bold = True
# 
#     # Установка полей для всего документа
#     sections = doc.sections
#     for section in sections:
#         section.top_margin = Cm(2)  # верхнее поле
#         section.bottom_margin = Cm(2)  # нижнее поле
#         section.left_margin = Cm(2)  # левое поле
#         section.right_margin = Cm(1)  # правое поле
# 
#     heading = doc.add_heading('ДОГОВОР КУПЛИ-ПРОДАЖИ', level=1)
#     heading.style = doc.styles['MyMainHeadingStyle']
#     heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading.paragraph_format.space_before = 0
#     heading.paragraph_format.space_after = 0
# 
#     heading_subMain = doc.add_heading('турпродукта', level=1)
#     heading_subMain.style = doc.styles['MyHeadingStyle']
#     heading_subMain.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_subMain.paragraph_format.space_before = 0
#     heading_subMain.paragraph_format.space_after = 0
# 
#     paragraph_about = doc.add_paragraph(
#         # f'\n{".".join((str(datetime.datetime.now().date()).split("-")[::-1]))}\n')
#         f'\n{booking.date_pay.date()}\n')
#     paragraph_about.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
#     paragraph_about.paragraph_format.space_after = 0
# 
#     paragraph_predislovie = doc.add_paragraph(
# f'''Туристическая компания ООО "Альбатрос", в лице Директора Байрамовой Эльвиры Рафкатовны действующего (ей) на основании устава, в дальнейшем именуемое «Агентство», с одной стороны, и {client.surname} {client.name} {client.patronymic} в дальнейшем именуемый(-ая) « КЛИЕНТ», с другой стороны, заключили договор о нижеследующем:
# '''
#     )
#     paragraph_predislovie.paragraph_format.space_after = 0
# 
#     heading_1 = doc.add_heading('1.ПРЕДМЕТ ДОГОВОРА.', level=1)
#     heading_1.style = doc.styles['MyHeadingStyle']
#     heading_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_1.paragraph_format.space_before = 0
#     heading_1.paragraph_format.space_after = 0
# 
#     paragraph_1 = doc.add_paragraph(
# f'''1.1. Агентство по поручению и за счет Клиента обязуется произвести бронирование туристского продукта в порядке и на условиях предусмотренных настоящим договором, Заявке на бронирование (Приложение №1 к договору) ,а Клиент обязуется произвести оплату туристского продукта на оказание услуг.
# '''
#     )
#     paragraph_1.paragraph_format.space_after = 0
# 
#     heading_2 = doc.add_heading(f'2 СВЕДЕНИЯ О ТУРОПЕРАТОРЕ.', level=1)
#     heading_2.style = doc.styles['MyHeadingStyle']
#     heading_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_2.paragraph_format.space_before = 0
#     heading_2.paragraph_format.space_after = 0
# 
#     paragraph_2 = doc.add_paragraph(
# f'''2.1 Туроператором, являющимся непосредственным исполнителем туристических услуг, входящих в тур. продукт, является юридическое лицо, указанное в Приложении № 2 
# 2.2 При заключении договора на бронирование отдельно авиабилетов, трансферов или отелей вне пакетного тура, Приложение № 2 не прилагается.
# '''
#     )
#     paragraph_2.paragraph_format.space_after = 0
# 
#     heading_3 = doc.add_heading(f'3 ОБЯЗАННОСТИ АГЕНТСТВА', level=1)
#     heading_3.style = doc.styles['MyHeadingStyle']
#     heading_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_3.paragraph_format.space_before = 0
#     heading_3.paragraph_format.space_after = 0
# 
#     paragraph_3 = doc.add_paragraph(
# f'''3.1.Агентство обязуется за вознаграждение, от своего имени, совершить по поручению и за счет Клиента юридические и иные действия, направленные на подбор, бронирование и оплату туристского продукта, потребительские свойства которого указаны в Заявке на бронирование (Приложение №1 к договору), являющейся неотъемлемой частью настоящего Договора.. Везде, где по тексту договора указан Клиент, имеются в виду также третьи лица, в интересах которых действует Клиент, сопровождающие его (сопровождаемые им) лица, в том числе несовершеннолетние, или иной заказчик туристского продукта.
# Турагент _______________ \t\t\t\t\t\t Клиент _______________'''
#     )
#     paragraph_3.paragraph_format.space_after = 0
# 
#     # Добавляем разрыв страницы
#     doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# 
#     paragraph_3_1 = doc.add_paragraph(
# f'''3.2.Туристский продукт, соответствующий характеристикам, указанным в Заявке на бронирование, формируется Туроператором, сведения о котором содержатся в Приложении № 2, являющимся неотъемлемой частью настоящего Договора.. Туроператор является лицом (исполнителем), обеспечивающим оказание Клиенту услуг, входящих в туристский продукт, и несет перед Клиентом ответственность за неоказание или ненадлежащее оказание Клиенту услуг, входящих в туристский продукт, независимо от того, кем должны были оказываться или оказывались эти услуги.
# 3.3.Агентство предоставляет Клиенту достоверные сведения о составе и характеристиках услуг, входящих в туристский продукт. Услуги, входящие в туристский продукт, непосредственно оказываются Туристу третьими лицами - туроператором, перевозчиком, отелем или иным средством размещения, страховщиком и прочими лицами, предоставляющими услуги, входящие в туристский продукт.
# 3.4. За оказанные Агентством Клиенту услуги по подбору, бронированию и оплате туристского продукта согласно п.п. 1.1. настоящего Договора Клиент выплачивает Агентству вознаграждение в виде разницы между полученными денежными средствами от Клиента в оплату заказанных услуг и денежными средствами, оплаченными поставщику услуг (Туроператору). Вознаграждение самостоятельно удерживается Агентством из всей суммы, полученной от Клиента.
# 3.5. Лицом, (исполнителем), оказывающим Клиенту услуги по договору о реализации туристского продукта, является туроператор
# '''
#     )
#     paragraph_3_1.paragraph_format.space_after = 0
# 
#     heading_4 = doc.add_heading(f'4 ОБЯЗАТЕЛЬСТВА КЛИЕНТА', level=1)
#     heading_4.style = doc.styles['MyHeadingStyle']
#     heading_4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_4.paragraph_format.space_before = 0
#     heading_4.paragraph_format.space_after = 0
# 
#     paragraph_4 = doc.add_paragraph(
# f'''Клиент обязуется:
# 4.1.Произвести предоплату, но не менее 50% от общей стоимости тура. Курс валюты на момент заключения договора может меняться, а доплата по туру пересчитывается по курсу сайта Туроператора на день доплаты .
# 4.2. Оплатить полную стоимость тура в течение 5 банковских дней с момента его подтверждения, но не позднее, чем за 28 суток до предполагаемой даты отъезда, если иное не указано в Заявке на бронирование (Приложение №1 к договору);
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_4.paragraph_format.space_after = 0
# 
#     # Добавляем разрыв страницы
#     doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# 
#     paragraph_4_1 = doc.add_paragraph(
# f'''4.3. Предоставить в Агентство действительный заграничный паспорт и все необходимые достоверные документы для оформления визы не позднее, чем за 20 суток до начала тура, а также предоставить Агентству полную и достоверную информацию об адресе фактического места жительства, контактные телефоны для оперативной связи в момент заключения настоящего Договора;
# 4.4. Прибыть в аэропорт на регистрацию не позднее чем, за 3 часа до вылета самолета (поезда, парома, автобуса) для самостоятельного прохождения пограничного и таможенного контроля; уточнить за 24 часа до вылета (выезда) у Агентства или в справочной аэропорта вылета время и место вылета рейса, причем получение такой информации является обязанностью Клиента.
# 4.5. Соблюдать пограничные и таможенные правила РФ и посещаемой страны, правила авиакомпании по провозу багажа;
# 4.6. Оплатить до выезда из отеля счета за пользование мини-баром в номере, телефонные переговоры, и другие дополнительные услуги отеля;
# 4.7. Соблюдать во время путешествия правила личной безопасности; законодательство страны временного пребывания, уважать ее социальное устройство обычаи, традиции, религиозные верования; правила въезда (выезда) в страну временного пребывания; рекомендации касающиеся профилактики инфекционных и паразитарных заболеваний; правила пользования отдельными услугами отеля (например, водные горки и аттракционы, spa-салонов и т.п.).
# 4.8. В случае расторжения Договора по инициативе Клиента он обязан оплатить все фактически понесенные расходы Агентства, указанные в Договоре на сайте Туроператора ,связанные с исполнением условий настоящего Договора (п.5.2)
# 4.9. Клиент берет на себя всю ответственность, включая финансовую, за любые совершенные им действия или решения, принимаемые в ходе поездки, в т.ч. самостоятельно вносимые изменения во время путешествия, или опоздания на свои рейсы, трансфер, выход к трансферу, выезд из отеля, поведение во время рейса, трансфера, экскурсий и т.п., а также несет ответственность за соблюдение законодательства страны пребывания.
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_4_1.paragraph_format.space_after = 0
# 
#     # Добавляем разрыв страницы
#     doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# 
#     paragraph_4_2 = doc.add_paragraph(
# f'''4.10. Не позднее, чем за 1 день до вылета получить полный пакет документов для осуществления тура (в т.ч. билеты, ваучер, медстраховку и тд.) в офисе Агентства, либо по электронной почте, указанной в настоящем Договоре. По согласованию с Агентством, Клиент может получить пакет документов на стойке Туроператора в аэропорту за 2,5 часа до вылета. Клиент берет на себя всю ответственность, в том числе и финансовую, за действия или (бездействие) приведшие к невозможности получения им документов на тур.
# '''
#     )
#     paragraph_4_2.paragraph_format.space_after = 0
# 
#     heading_5 = doc.add_heading(f'5 ОСОБЫЕ УСЛОВИЯ', level=1)
#     heading_5.style = doc.styles['MyHeadingStyle']
#     heading_5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_5.paragraph_format.space_before = 0
#     heading_5.paragraph_format.space_after = 0
# 
#     paragraph_5 = doc.add_paragraph(
# f'''5.1. Обязательства туроператора перед Клиентом считаются возникшими с момента подтверждения туроператором возможности оказать туристские услуги. Подтверждение заказанных услуг означает, что эти услуги будут оказаны в порядке и на условиях, изложенных в настоящем договоре. Срок подтверждения возможности оказания туристских услуг устанавливается в три рабочих дня с момента подписания договора. Заказанные услуги Клиента являются безотзывной офертой на срок в три рабочих дня. Подтверждение заказанных услуг является акцептом. В случае не подтверждения заказанных Клиентом услуг, настоящий Договор считается незаключенным. В этом случае денежные средства, внесенные Клиентом в Агентство, возвращаются ему в полном объеме. Получение информации о подтверждении (либо не подтверждении) возможности оказать туристские услуги является обязанностью Клиента. Обязательства Агентства перед Клиентом считаются выполненными с момента полной оплаты Туроператору подтвержденных им (Туроператором) туристских услуг, указанных в Приложении №1 настоящего Договора.
# 5.2. Клиент предупрежден, что при оказании визовой поддержки, в тех случаях, когда это необходимо, консульства могут потребовать для оформления визы предоставление доказательств намерения Клиента посетить страну временного пребывания (авиабилет, страховку, ваучер и т.п.). Клиент предупрежден, что при подаче документов на визу в срок меньше установленного консульством, своевременное получение визы не гарантируется и Агентство не несет за это ответственности, и Клиент согласен относить в данном случае все фактически понесенные расходы Агентства (см п.5.2) на свой счет.
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_5.paragraph_format.space_after = 0
# 
#     # Добавляем разрыв страницы
#     doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# 
#     paragraph_5_1 = doc.add_paragraph(
# f'''Подписывая настоящий Договор, Клиент подтверждает, что проинформирован о сроках получения визы, и, в случае если виза не будет получена, согласен относить в данном случае все фактически понесенные расходы Агентства на свой счет.
# В случае несвоевременного предоставления Клиентом документов на оформление въездных виз, Агентство производит аннуляцию тура. В этом случае
# Клиент обязуется компенсировать Агентству фактически понесенные им расходы, вызванные аннуляцией тура в порядке и на условиях, предусмотренных настоящим Договором.
# В случае, если Клиент отказывается от оформленной визы, услуга по оформлению визы считается выполненной Агентством, поэтому консульский сбор не возвращается. В этом случае, Клиент обязуется написать заявление на отказ от визы, после чего паспорта с оформленными визами возвращаются в консульство для аннулирования визы.
# Отказ в выдаче въездной визы консульством иностранного государства, а также просрочка ее выдачи, что влечет невозможность предоставления туристского продукта Клиенту, не являются форс-мажорным обстоятельством. Любые фактические расходы, понесенные Агентством и связанные с таким отказом или просрочкой своевременной выдачи визы, в том числе выезд на собеседование в консульство, что повлекло невозможность воспользоваться туристским продуктом, относятся на счет Клиента, и возмещаются Клиентом Агентству в порядке и на условиях, предусмотренных настоящим Договором.
# 5.3. Перевозка Клиента по настоящему Договору выполняется по правилам перевозчика (авиакомпании и пр.). Билет Клиента является договором с перевозчиком, в соответствии, с которым всю ответственность за перевозку несет перевозчик по правилам перевозчика. Клиент предупрежден, что стоимость чартерных авиабилетов и авиабилетов на регулярные рейсы с невозвратным тарифом не возвращается независимо от срока отказа от тура.
# Подписывая настоящий Договор, Клиент подтверждает, что проинформирован и предупрежден о том, что при приобретении туристского продукта, в котором включена авиаперевозка, авиаперевозка вне зависимости от предлагаемой авиакомпании, является чартерной. Авиабилеты, приобретаемые при формировании туристского продукта по чартерному тарифу, являются невозвратными вне зависимости от времени отказа от авиаперевозки. 
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_5_1.paragraph_format.space_after = 0
# 
#     paragraph_5_2 = doc.add_paragraph(
# f'''Клиент, оплачивая стоимость авиабилета, уведомлен и принимает все условия авиаперевозки, в том числе условия договора воздушной перевозки между авиакомпанией и туроператором при формировании туристского продукта. Клиент принимает условия указанные в настоящем абзаце и осведомлен, что оплата перевозки и правила возврата денежных средств за авиабилеты согласно настоящему абзацу отличаются от правил, указанных в Воздушном кодексе РФ.
# Подписывая настоящий Договор, клиент подтверждает, что проинформирован и предупрежден о том, что Туроператор вправе без согласования с
# Агентством или Клиентом заменить авиакомпанию, тип воздушного судна, рейс, время и аэропорт вылета (прилета). Авиа (ж/д) рассадка при бронировании тура является предварительной, что не гарантирует ее сохранения к моменту вылета, и окончательно определяется Туроператором за 1 день до вылета, о чем
# Клиент оповещается по электронной почте или телефону. Уточнение окончательной рассадки и времени вылета является обязанностью Клиента.
# 5.4. Незнание Клиентом законов или обычаев страны пребывания не освобождает его от ответственности при их нарушении. Гид или сопровождающий не является комментатором закона и не разделяет ответственность по чужому действию или бездействию.
# 5.5. Туроператор оставляет за собой право, в случае необходимости, заменить отель, подтвержденный ранее, на отель той же либо более высокой категории.
# 5.6. Стороны по настоящему Договору договорились, что при нарушении п. 3.2. настоящего договора договор считается расторгнутым по инициативе
# Клиента, при этом Клиент обязан возместить Агентству фактически понесенные им расходы в соответствие настоящим Договора.
# 5.7. Стороны пришли к соглашению, что в случае не полной или несвоевременной оплаты Клиентом туристского продукта в соответствии с п.3.2 настоящего договора, либо непредвиденном введении новых или повышении действующих налогов и сборов, и (или) при резком изменении курса национальных валют,
# Агентство вправе изменить условия оплаты по настоящему Договору с учетом вышеизложенных изменений. Агентство выставляет Клиенту новый счет, в случае неоплаты которого, в указанные в счете сроки, Договор считается расторгнутым по инициативе Клиента. В случае расторжения настоящего Договора в соответствие с настоящим пунктом Клиент должен возместить Агентству фактически понесенных им расходов, связанных с исполнением Договора.
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_5_2.paragraph_format.space_after = 0
# 
#     paragraph_5_3 = doc.add_paragraph(
# f'''5.8. Подписывая настоящий Договор, Клиент подтверждает, что в соответствие со ст. 10 Федерального закона «Об основах туристской деятельности в
# Российской Федерации» полностью проинформирован о потребительских свойствах туристского продукта - о программе пребывания, маршруте и об условиях путешествия, включая информацию о средствах размещения, об условиях проживания (месте нахождения средства размещения, его категории) и питания, услугах по перевозке туриста в стране (месте) временного пребывания, о наличии экскурсовода (гида), гида-переводчика, инструктора-проводника, а также о дополнительных услугах. Клиент имеет право внести в настоящий Договор и иные условия, которые он считает существенными при заказе конкретного туристского продукта.
# 5.9. Подписывая настоящий Договор, Клиент подтверждает, что уведомлен о том, что в соответствии с действующим законодательством РФ, страховой полис, является договором на предоставление медицинских услуг и возмещение расходов связанных в предоставлением медицинской помощи между страховой компанией и Клиентом выезжающим за рубеж. Все условия страхования указаны в получаемом Клиентом полисе и правилах страхования. Убытки и другой любой ущерб, нанесенный здоровью и/или имуществу Клиента, понесенный последним в связи с неисполнением или ненадлежащим исполнением страховой компанией обязательств по заключенному договору страхования, подлежит возмещению страховой компанией; доказательством факта заключения договора между Клиентом и страховой компанией является страховой полис, переданный Клиенту. В связи с этим, все заявления, претензии, связанные с наступлением страхового случая, неисполнением или ненадлежащим исполнением страховой компанией принятых на себя обязательств по договору страхования, предъявляются Клиентом непосредственно в страховую компанию, полис которой был выдан Клиенту. Страховой полис и любые иные документы, подтверждающие наступление страхового случая и размер понесенных расходов в связи с наступлением страхового случая, необходимо сохранить до предъявления требований в страховую компанию. Одно из основных обязательств Застрахованного лица при наступлении страхового случая - незамедлительное уведомление об этом страховой компании и неуклонное следование ее указаниям (средства связи со страховой компанией указаны в страховом полисе). Следует иметь в виду, что такое уведомление, как правило, происходит по телефону, поэтому если застрахованный находится в отеле и звонит из номера, то отель выставит счет за услуги международной телефонной связи, который обязана, будет оплатить страховая компания.
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_5_3.paragraph_format.space_after = 0
# 
#     paragraph_5_4 = doc.add_paragraph(
# f'''5.10. Подписывая настоящий Договор, Клиент подтверждает, что у него есть или будут до момента начала тура все необходимые документы для совершения путешествия как его самого, так и лиц, в отношении которых заказывается туристских продукт, в т.ч. имеются загранпаспорта с необходимым сроком действия; разрешающие документы от законных представителей детей; отсутствуют какие-либо противопоказания для того, чтобы воспользоваться туристским продуктом как по медицинским так и иным основаниям и т.п. Клиент подтверждает, что все убытки, материальный и моральный вред, вызванный невозможностью воспользоваться туристским продуктом из-за подлинности или действительности документов, указанных в настоящем пункте, полностью относится на его счет. С целью недопущения возникновения проблем при прохождении российской и иностранной границы Клиент должен до подписания настоящего Договора получить консультации по этим документам и условиям выезда из Российской Федерации и обратного въезда в соответствующем подразделении Федерального миграционной службы.
# 5.11. Подписывая настоящий Договор, Клиент подтверждает, что проинформирован и предупрежден о том, что с целью недопущения у него убытков, связанных с возмещением фактически понесенных расходов Агентства из-за действий консульских учреждений иностранных государств или невозможностью совершения путешествия по независящим от Клиента причинам он может застраховать расходы, связанные с отменой поездки за границу или изменением сроков пребывания за границей. Услуга по оформлению страховки от невыезда считается исполненной с момента подтверждения туроператором заказанных услуг. В случае отсутствия такой страховки все убытки, связанные с возмещением фактически понесенных расходов относятся на счет Клиента. Наличие у Клиента указанной в настоящем пункте страховки подтверждает, что он был полностью проинформирован об условиях страхования.
# 5.12. Подписывая настоящий Договор, Клиент подтверждает, что проинформирован и предупрежден о том, что если невозможность воспользоваться туристским продуктом будет вызвана виной Клиента, то услуги подлежат оплате в полном объеме. Виной Клиента признается любое действие или бездействие, приведшее к невозможности воспользоваться туристским продуктом, и которое не связано с действиями (бездействием) Агентства или Туроператора в виде нарушения своих обязательств по настоящему Договору, а также действиями лиц, за которые ни одна из сторон не отвечает.
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_5_4.paragraph_format.space_after = 0
# 
#     # Добавляем разрыв страницы
#     doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# 
#     paragraph_5_5 = doc.add_paragraph(
# f'''5.13. Подписывая настоящий Договор, Клиент дает свое согласие на использование его персональных данных, которые по решению Клиента как субъекта персональных данных являются общедоступными, и могут быть использованы, как Агентством, так и третьими лицами ( в т.ч. туроператором, принимающей стороной, транспортными компаниями, авиаперевозчиками и др) в том объеме, который будет необходим Агентсту и неопределенному кругу лиц. Клиент подтверждает, что действует от всех учестников тура, имея от них согласие на использование их персональных данных. Все персональные данные предоставленные Агентству являются общедоступными, соблюдение конфиденциальности при их использовании не требуется.
# 5.14. Подписывая настоящий Договор, Клиент подтверждает, что полностью проинформирован об обстоятельствах, указанных в ст. 14 Федерального закона «Об основах туристской деятельности в Российской Федерации» (в т.ч. о правилах въезда в страну (место) временного пребывания и выезда из страны (места) временного пребывания, наличие необходимых документов для совершения путешествия; о возможных опасностях, с которыми Клиент может встретиться при совершении путешествия; о таможенных, пограничных, медицинских, санитарно-эпидемиологических и иных правилах (в объеме, необходимом для совершения путешествия); о реквизитах дипломатических представительств и консульских учреждений Российской Федерации; об адресе (месте пребывания) и номере контактного телефона в стране (месте) временного пребывания руководителя группы несовершеннолетних туристов в случае, если туристский продукт включает в себя организованный выезд группы несовершеннолетних туристов без сопровождения родителей, усыновителей, опекунов или попечителей; о национальных и религиозных особенностях страны (места) временного пребывания; об иных особенностях путешествия
# '''
#     )
#     paragraph_5_5.paragraph_format.space_after = 0
# 
#     heading_6 = doc.add_heading(f'6 ОТВЕТСТВЕННОСТЬ СТОРОН:', level=1)
#     heading_6.style = doc.styles['MyHeadingStyle']
#     heading_6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_6.paragraph_format.space_before = 0
#     heading_6.paragraph_format.space_after = 0
# 
#     paragraph_6 = doc.add_paragraph(
# f'''6.1. Агентство несет ответственность только в рамках своих договорных обязательств согласно п. 1.1 настоящего договора при условии полной оплаты тура в установленные сроки.
# Турагент _______________ \t\t\t\t\t\t Клиент _______________ 
# '''
#     )
#     paragraph_6.paragraph_format.space_after = 0
# 
#     # Добавляем разрыв страницы
#     doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# 
#     paragraph_6_1 = doc.add_paragraph(
# f'''6.2. Клиент имеет право отказаться от подтвержденного тура в любой момент, при этом возместив Агентству фактически понесенные расходы.
# Стороны пришли к соглашению о том, что фактически понесенными расходами они будут считать любые документально подтвержденные затраты осуществленные (понесенные) Агентством, оформленными в соответствии с законодательством
# Российской Федерации, либо документами, оформленными в соответствии с обычаями делового оборота, применяемыми в иностранном государстве, на территории которого были произведены соответствующие расходы, связанные с исполнением настоящего Договора. Фактически понесенными расходами являются также расходы Агентства, связанные с уплатой
# Туроператору или иным контрагентам штрафных санкций (пеней, неустоек и т.п), которые Агентство вынуждено понести в связи с отказом Клиента от
# туристского продукта по любой причине, либо по причине невыполнения Клиентом своих обязательств по настоящему Договору, либо в результате аннуляции тура и (или) расторжения договора по требованию Агента, либо в результате невозможности предоставления Клиенту туристского продукта по причинам, за которые ни одна из сторон не отвечает. Фактически понесенные расходы считаются связанными с исполнением настоящего договора, если они связаны с любыми действиями (бездействими) Агентства по выполнению своих обязательств по настоящему Договору, в т.ч. по аннуляции заказа у
# туроператора и расторжении настоящего Договора по инициативе Агентства и в остальных оговоренных настоящим Договором случаях.
# Фактически понесенные расходы признаются произведенными независимо от даты их совершения, например, они могут быть произведены как до, так и после расторжения или изменения договора, как до так и после получения требования одной из сторон о расторжении или изменении договора, как до так и после возникновения условий препятствующих совершению путешествия, и т.п. Подписывая настоящий Договор, Клиент подтверждает, что проинформирован и ознакомлен с условиями договора, заключенного между Агентством и туроператором, по которому будет осуществляться бронирование тура Клиента. Клиент предупрежден обо всех условиях штрафных санкций туроператора, у которого забронирован туристский продукт, и готов относить их на свой счет.
# При отказе от тура по уважительным причинам Агентство предпримет все меры к минимизации расходов клиента.
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_6_1.paragraph_format.space_after = 0
# 
#     paragraph_6_2 = doc.add_paragraph(
# f'''6.3. Агентство не несет ответственность перед Клиентом, не возвращает полную или частичную стоимость тура и компенсацию за моральный ущерб при:
# 6.3.1. нарушении Клиентом положений настоящего договора, и при невыполнении Клиентом обязательств, изложенных в статье 3 настоящего договора;
# 6.3.2. отказе иностранного государства в выдаче въездных виз Клиенту по маршруту Тура, за задержки при рассмотрении документов консульством иностранного государства
# 6.3.3. не прохождении Клиентом таможенного, санитарного, пограничного контроля и других служб аэропортов, в том числе если это связанно с неправильным оформлением или недействительностью паспорта Клиента, либо отсутствием записи о членах семьи в паспорте клиента или отсутствием или неправильным оформлением доверенностей на несовершеннолетних;
# 6.3.4. изменении ценовой политики авиакомпаний, изменения тарифов на забронированные авиабилеты, за задержку вылетов и прилетов, замену типа самолета, отмену рейсов, за доставку и сохранность багажа Клиентов;
# 6.3.5. возникновении проблем, трудностей и последствий, возникающих у Клиента при утере Клиентом загранпаспорта или других документов необходимых для осуществления тура;
# 6.3.6. самостоятельном изменении Клиентом отдельных элементов программы (экскурсионной программы, трансфера, места и уровня проживания, несвоевременной явки к месту сбора группы и др.), вызвавших дополнительные затраты со стороны Клиента;
# 6.3.7. несоответствии предоставленных услуг, необоснованным ожиданиям Клиента и его субъективной оценке;
# 6.3.8. возникновении проблем, связанных с сохранностью личного багажа, ценностей и документов Клиента в течение всего периода поездки. Агентство предупреждает Клиента о необходимости принятия собственных мер, направленных на обеспечение сохранности личных вещей, ценностей и документов на всем протяжении поездки;
# 6.3.9. возникновении проблем, связанных с подлинностью документов, предоставляемых Клиентом для оформления и организации туристической поездки (паспорт, справка, доверенность и т.д.) и не несет ответственности за возможные последствия, связанные с этими обстоятельствами.
# 6.3.10. при изменении администрацией отеля формулы питания, перечня бесплатных/платных услуг отеля, невозможности предоставления доступа в интернет и пр., независящего от Агентства.
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# 6.3.11. при осуществлении в отеле либо на близлежащей территории строительных, ремонтных или иных восстановительных работ, о проведении которых не было известно Агентству.
# '''
#     )
#     paragraph_6_2.paragraph_format.space_after = 0
# 
#     paragraph_6_3 = doc.add_paragraph(
# f'''6.4 Агентство предоставляет информацию по турам от разных туроператоров и не несет ответственности за выбор туроператора Клиентом. Агентство не имеет полномочий по оценке финансового положения туроператора. Таким образом, Клиент сам делает вывод о надежности того или иного туроператора.
# '''
#     )
#     paragraph_6_3.paragraph_format.space_after = 0
# 
#     heading_7 = doc.add_heading(f'7 ПРОЧИЕ УСЛОВИЯ', level=1)
#     heading_7.style = doc.styles['MyHeadingStyle']
#     heading_7.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_7.paragraph_format.space_before = 0
#     heading_7.paragraph_format.space_after = 0
# 
#     paragraph_7 = doc.add_paragraph(
# f'''7.1 Ответственность перед Клиентом за неоказание или ненадлежащее оказание услуг, входящих в туристский продукт, отвечающий указанным в Заявке на бронирование требованиям Клиента, независимо от того, кем должны были оказываться или оказывались эти услуги, несет Туроператор, сведения о котором содержатся в Приложении № 2 к настоящему Договору. Туроператор несет ответственность перед Клиентом за неисполнение или ненадлежащее исполнение обязательств по Договору о реализации туристского продукта, заключенному турагентом как от имени туроператора, так и от своего имени.
# 7.2 При наличии каких-либо замечаний относительно качества услуг, оказываемых на протяжении путешествия и перечисленных в Заявке, или замечаний относительно действий третьих лиц, непосредственно оказывающих Клиенту услуги, Агентство рекомендует Клиенту незамедлительно обратиться к Туроператору и представителям принимающей стороны на местах, по телефонам, указанным в Договоре и приложениях к нему, а также в ваучере, программе пребывания и памятке.
# 7.3. Настоящим договором предусматривается претензионный порядок разрешения споров. При возникновении у Клиента претензий к качеству туристского продукта, Клиент обязан сообщить об этом представителю туроператора через гида. При невозможности разрешить проблему на месте гид составляет протокол, заверенный гидом, который служит основанием для получения компенсации от отеля или виновного лица. Претензии по качеству туристского продукта от Клиента принимаются в течение 20 (двадцати) календарных дней со дня окончания тура с приложением всех документов, подтверждающих не предоставление или некачественное предоставление туристских услуг, и подлежат рассмотрению в течение 10 (десяти) дней со дня получения претензий. 
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_7.paragraph_format.space_after = 0
# 
#     paragraph_7_1 = doc.add_paragraph(
# f'''В случаях неисполнения или ненадлежащего исполнения обязательств по оказанию Клиенту услуг, входящих в туристский продукт по настоящему Договору, при наличии оснований для уплаты страхового возмещения по договору страхования ответственности Клиент вправе в пределах суммы финансового обеспечения предъявить письменное требование об уплате страхового возмещения непосредственно страховщику - организации, предоставившей финансовое обеспечение и указанной в Приложении №2 настоящего Договора. Письменное требование Клиента об уплате страхового возмещения по договору страхования ответственности должно быть предъявлено страховщику в течение срока действия договора страхования должно быть предъявлено страховщику в течение срока действия договора страхования. Основанием для уплаты денежной суммы по договору страхования ответственности является факт установления обязанности возместить Клиенту реальный ущерб, возникший в результате неисполнения или ненадлежащего исполнения обязательств, если это является существенным нарушением условий договора. Право требования денежной компенсации у Клиента от гаранта, выдавшего банковскую гарантию, гарантирующего финансовое обеспечение, возникает после вступления в законную силу решения суда по установлению факта нарушения прав Клиента, в случае отказа Туроператора от исполнения данного судебного решения.
# 7.4. Все споры между сторонами по настоящему договору, не урегулированные путем переговоров, передаются на рассмотрение судебных органов по месту нахождения ответчика.
# 7.5. Договор вступает в действие с момента подписания и действует до момента передачи Клиенту документов на тур. Стороны договорились, что в рамках настоящего Договора допускается передача документов на тур путем их отправки по средствам электронной почты, указанной в реквизитах Сторон.
# 7.6. Настоящий договор составлен в двух экземплярах на русском языке и хранится по одному у каждой из сторон. Ничтожность или недействительность отдельного пункта настоящего Договора не означает автоматической ничтожности или недействительности других пунктов или Договора в целом.
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_7_1.paragraph_format.space_after = 0
# 
#     # Добавляем разрыв страницы
#     doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# 
#     paragraph_7_2 = doc.add_paragraph(
# f'''7.7. Любые изменения и дополнения к настоящему Договору действительны, если они заключены в письменной форме или отправлены по средствам электронной почты, указанной в реквизитах сторон настоящего Договора. Стороны признают юридическую силу за текстами электронных писем и дoкyмeнтaми, направленными по электронной почте (e-mail), и признают их равнозначными дoкyмeнтaм на бумажных носителях, подписанным собственноручной подписью, т.к. только сами Стороны и уполномоченные ими лица имеют доступ к соответствующим средствам связи - адресам электронной почты, указанным в Договоре в реквизитах Сторон. Доступ к электронной почте каждая Сторона осуществляет по паролю и обязуется сохранять его конфиденциальность. Все уведомления и сообщения, отправленные Сторонами друг другу по указанным в реквизитах Сторон адресам электронной почты и/или по телефонным номерам, признаются Сторонами официальной перепиской в рамках настоящего Договора. Датой передачи соответствующего сообщения считается день отправления сообщения электронной почты или СМС на указанный номер телефона. Ответственность за получение сообщений и уведомлений вышеуказанным способом лежит на получающей Стороне.
# 7.8. Во всем том, что не урегулировано настоящим договором стороны руководствуются действующим законодательством Российской Федерации. Приложения к настоящему договору:
# 1 Заявка на бронирование. ПРИЛОЖЕНИЕ №1
# 2 Сведения о Туроператоре. ПРИЛОЖЕНИЕ № 2
# Турагент _______________ \t\t\t\t\t\t Клиент _______________
# '''
#     )
#     paragraph_7_2.paragraph_format.space_after = 0
# 
#     # Добавляем разрыв страницы
#     doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# 
#     heading_8 = doc.add_heading(f'ПРИЛОЖЕНИЕ 1', level=1)
#     heading_8.style = doc.styles['MyMainHeadingStyle']
#     heading_8.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_8.paragraph_format.space_before = 0
#     heading_8.paragraph_format.space_after = 0
# 
#     heading_8_1 = doc.add_heading(f'Заявка на бронирование\n', level=1)
#     heading_8_1.style = doc.styles['MyHeadingStyle']
#     heading_8_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_8_1.paragraph_format.space_before = 0
#     heading_8_1.paragraph_format.space_after = 0
# 
#     paragraph_8 = doc.add_paragraph(
# f'''Идентификатор бронирования: {booking.id}
# Страна пребывания: {tour.country_stay}
# Пункт отправления: {tour.city_departure}
# Пункт пребывания: {tour.city_stay}
# Дата пребывания: {tour.date_departure.date()}
# Дата пребывания: {tour.date_return.date()}
# Цена: {tour.price}₽
# '''
#     )
#     paragraph_8.paragraph_format.space_after = 0
# 
#     # Добавляем разрыв страницы
#     doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
# 
#     heading_9 = doc.add_heading(f'ПРИЛОЖЕНИЕ 2', level=1)
#     heading_9.style = doc.styles['MyMainHeadingStyle']
#     heading_9.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_9.paragraph_format.space_before = 0
#     heading_9.paragraph_format.space_after = 0
# 
#     heading_9_1 = doc.add_heading(f'Сведения о Туроператоре\n', level=1)
#     heading_9_1.style = doc.styles['MyHeadingStyle']
#     heading_9_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     heading_9_1.paragraph_format.space_before = 0
#     heading_9_1.paragraph_format.space_after = 0
# 
#     paragraph_9 = doc.add_paragraph(
# f'''Название: {to.title}
# Контакты
# Электронная почта: {cotactTO.email}
# Телефон: {cotactTO.tel_number}
# '''
#     )
#     paragraph_9.paragraph_format.space_after = 0
# 
#     # Получаем нижний колонтитул для этой секции
#     footer = doc.sections[0].footer
# 
#     # Создаем абзац для добавления нумерации страниц
#     paragraph_col = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
#     paragraph_col.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# 
#     # Создаем XML-элемент с нумерацией страниц
#     page_num = OxmlElement('w:fldSimple')
#     page_num.set(qn('w:instr'), 'PAGE \* MERGEFORMAT')
#     paragraph_col._element.append(page_num)
# 
#     temp_file = io.BytesIO()
#     doc.save(temp_file)
#     temp_file.seek(0)
# 
#     return send_file(temp_file, as_attachment=True, download_name='generated_file.docx')
